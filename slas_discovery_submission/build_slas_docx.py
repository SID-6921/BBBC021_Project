from __future__ import annotations

from pathlib import Path
import json

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = Path(r"C:\Users\nanda\Downloads\Manuscript (2).docx")
OUTPUT_PATH = ROOT / "slas_discovery_submission" / "BBBC021_SLAS_Discovery_Submission.docx"
MARKDOWN_PATH = ROOT / "slas_discovery_submission" / "manuscript_slas_discovery.md"

# Main text figures (8 maximum for SLAS Discovery).
MAIN_FIGURES = {
    "figure1_pipeline_workflow.png":
        "Figure 1. End-to-end analysis workflow for BBBC021 data curation, preprocessing, detection, "
        "feature extraction, modeling, and validation.",
    "figure2_detection_samples.png":
        "Figure 2. Representative detection overlays highlighting bright-region contour localization "
        "across sample images.",
    "figure4_classification_roc.png":
        "Figure 3. ROC curves for feature-based supervised classifiers "
        "(Logistic Regression and Random Forest).",
    "figure4_feature_importance.png":
        "Figure 4. Random Forest feature-importance ranking.",
    "figure9_deep_learning_roc.png":
        "Figure 5. ROC curves for deep-learning models "
        "(CNN scratch, ResNet-18 scratch, ResNet-18 pretrained).",
    "figure14_calibration_curves_multiclass.png":
        "Figure 6. Reliability (calibration) curves for all five evaluated models.",
    "figure16_pca_batch_after.png":
        "Figure 7. PCA of feature vectors after within-batch z-score normalization, "
        "showing reduced batch-driven separation.",
    "figure17_feature_ablation.png":
        "Figure 8. Feature-ablation impact on Random Forest macro-OvR ROC AUC.",
}

# Supplementary figures moved out of the main text.
SUPPLEMENTARY_FIGURES = {
    "figure3_spot_count_boxplot.png":
        "Supplementary Figure S1. Group-wise distribution of spot count across MOA classes.",
    "figure3_mean_intensity_boxplot.png":
        "Supplementary Figure S2. Group-wise distribution of mean intensity across MOA classes.",
    "figure3_density_spots_per_10k_px_boxplot.png":
        "Supplementary Figure S3. Group-wise distribution of spot density across MOA classes.",
    "figure5_pca_clustering.png":
        "Supplementary Figure S4. PCA projection with MOA class overlay for unsupervised structure analysis.",
    "figure6_batch_robustness.png":
        "Supplementary Figure S5. Detection performance across weekly acquisition batches.",
    "figure7_threshold_sensitivity.png":
        "Supplementary Figure S6. Sensitivity of detection metrics across threshold configurations.",
    "figure8_robustness_classification.png":
        "Supplementary Figure S7. Classification performance under alternative detection threshold settings.",
    "figure9_deep_learning_training.png":
        "Supplementary Figure S8. Training and validation loss curves for the compact CNN.",
    "figure11_confusion_matrix_lr_12class.png":
        "Supplementary Figure S9. Confusion matrix for Logistic Regression on the 165-image "
        "held-out test set (10 MOA classes).",
    "figure12_confusion_matrix_rf_12class.png":
        "Supplementary Figure S10. Confusion matrix for Random Forest on the 165-image "
        "held-out test set (10 MOA classes).",
    "figure13_confusion_matrix_resnet_pretrained_12class.png":
        "Supplementary Figure S11. Confusion matrix for ImageNet pretrained ResNet-18 on the "
        "165-image held-out test set (10 MOA classes).",
    "figure15_pca_batch_before.png":
        "Supplementary Figure S12. PCA of feature vectors before within-batch z-score normalization, "
        "showing pronounced batch clustering.",
    "supplementary_figure_s1_feature_correlation.png":
        "Supplementary Figure S13. Spearman correlation matrix of top morphological features.",
}

TABLE_SPECS = [
    ("results_summary/advanced_model_metrics.csv",
     "Table 1. Model performance metrics with 95% bootstrap confidence intervals."),
    ("results_summary/model_comparison_transfer_learning.csv",
     "Table 2. Transfer learning comparison: CNN scratch vs ResNet-18 scratch vs ResNet-18 pretrained."),
    ("results_summary/advanced_calibration_ece.csv",
     "Table 3. Calibration error summary (ECE per model)."),
    ("results_summary/advanced_feature_ablation.csv",
     "Table 4. Feature-ablation outcomes (Random Forest macro-OvR AUC by features removed)."),
    ("results_summary/advanced_biological_validation.csv",
     "Table 5. Biological validation statistics for top morphological features."),
    ("results_summary/advanced_computational_cost.csv",
     "Table 6. Computational cost comparison (training time, peak RAM, pretrained weights)."),
    ("results_summary/advanced_nested_cv.csv",
     "Table 7. Nested cross-validation stability summary."),
    ("results_summary/robustness_classification.csv",
     "Table 8. Robustness classification across threshold settings."),
]


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_normal_font(doc: Document) -> None:
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(12)


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_dataframe_table(doc: Document, df: pd.DataFrame, caption: str) -> None:
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_json_table(doc: Document, data: dict, caption: str) -> None:
    rows = []
    for cmp_name, values in data.items():
        if isinstance(values, dict):
            row = {"comparison": cmp_name}
            row.update(values)
            rows.append(row)
    if rows:
        add_dataframe_table(doc, pd.DataFrame(rows), caption)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.2))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_markdown_content(doc: Document, md_text: str) -> None:
    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(16)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith(("1. ", "2. ", "3. ", "4. ", "5. ",
                               "6. ", "7. ", "8. ", "9. ")):
            doc.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip("*"))
            r.bold = True
        else:
            # Handle inline bold subheadings used in the structured abstract.
            if line.startswith("**") and ".**" in line:
                p = doc.add_paragraph()
                bold_end = line.index(".**") + 3
                r_bold = p.add_run(line[2:bold_end - 1])
                r_bold.bold = True
                remaining = line[bold_end:].strip()
                if remaining:
                    p.add_run(" " + remaining)
            else:
                doc.add_paragraph(line)


def main() -> None:
    if TEMPLATE_PATH.exists():
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    clear_document(doc)
    set_normal_font(doc)

    md = MARKDOWN_PATH.read_text(encoding="utf-8")
    add_markdown_content(doc, md)

    # Main text figures (8 figures).
    doc.add_heading("Figures", level=1)
    fig_dir = ROOT / "final_figures"
    for fig_name, caption in MAIN_FIGURES.items():
        fig_path = fig_dir / fig_name
        if fig_path.exists():
            add_figure(doc, fig_path, caption)
        else:
            print(f"WARNING: main figure not found: {fig_path}")

    # Tables.
    doc.add_heading("Tables", level=1)
    for rel, caption in TABLE_SPECS:
        p = ROOT / rel
        if p.exists():
            add_dataframe_table(doc, pd.read_csv(p), caption)
        else:
            print(f"WARNING: table CSV not found: {p}")

    # DeLong JSON table.
    delong = ROOT / "results_summary" / "advanced_delong_tests.json"
    if delong.exists():
        add_json_table(
            doc,
            json.loads(delong.read_text(encoding="utf-8")),
            "Table 9. DeLong test results (macro-OvR AUC) comparing deep-learning models "
            "against Random Forest.",
        )

    # Supplementary figures section.
    doc.add_page_break()
    doc.add_heading("Supplementary Material", level=1)
    for fig_name, caption in SUPPLEMENTARY_FIGURES.items():
        fig_path = fig_dir / fig_name
        if fig_path.exists():
            add_figure(doc, fig_path, caption)
        else:
            print(f"WARNING: supplementary figure not found: {fig_path}")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
