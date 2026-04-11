from __future__ import annotations

from pathlib import Path
import json

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = Path(r"C:\Users\nanda\Downloads\Manuscript (2).docx")
OUTPUT_PATH = ROOT / "results_summary" / "BBBC021_Manuscript_Submission_v4.docx"
MARKDOWN_PATH = ROOT / "results_summary" / "manuscript_mdpi_toxins_v3.md"

FIGURE_CAPTIONS = {
    "figure1_pipeline_workflow.png": "Figure 1. End-to-end analysis workflow for BBBC021 data curation, preprocessing, detection, feature extraction, modeling, and validation.",
    "figure2_detection_samples.png": "Figure 2. Representative detection overlays highlighting bright-region contour localization across sample images.",
    "figure3_spot_count_boxplot.png": "Figure 3A. Group-wise distribution of spot-count features.",
    "figure3_mean_intensity_boxplot.png": "Figure 3B. Group-wise distribution of mean-intensity features.",
    "figure3_density_spots_per_10k_px_boxplot.png": "Figure 3C. Group-wise distribution of spot-density features.",
    "figure4_classification_roc.png": "Figure 4A. ROC curves for feature-based supervised models.",
    "figure4_feature_importance.png": "Figure 4B. Random Forest feature-importance ranking.",
    "figure5_pca_clustering.png": "Figure 5. PCA projection with clustering overlay for unsupervised structure analysis.",
    "figure6_batch_robustness.png": "Figure 6. Detection robustness across weekly batches.",
    "figure7_threshold_sensitivity.png": "Figure 7. Sensitivity profile of metrics across threshold configurations.",
    "figure8_robustness_classification.png": "Figure 8. Classification performance across robustness configurations.",
    "figure9_deep_learning_roc.png": "Figure 9A. ROC curve for compact CNN baseline.",
    "figure9_deep_learning_training.png": "Figure 9B. CNN optimization curve (training vs validation loss).",
    "figure10_confusion_matrix_rf.png": "Figure 10. Confusion matrix for Random Forest on advanced test split.",
    "figure11_confusion_matrix_cnn.png": "Figure 11. Confusion matrix for compact CNN on advanced test split.",
    "figure11_confusion_matrix_lr_12class.png": "Figure 11. 10-class confusion matrix for Logistic Regression (reviewer revision).",
    "figure12_confusion_matrix_rf_12class.png": "Figure 12. 10-class confusion matrix for Random Forest (reviewer revision).",
    "figure13_confusion_matrix_resnet_pretrained_12class.png": "Figure 13. 10-class confusion matrix for pretrained ResNet-18 (reviewer revision).",
    "figure12_confusion_matrix_resnet18.png": "Figure S2. Confusion matrix for ResNet-18 scratch baseline.",
    "figure13_calibration_curves.png": "Figure 14. Reliability (calibration) curves for all evaluated models.",
    "figure14_pca_batch_before.png": "Figure 14. PCA visualization of batch effects before normalization.",
    "figure15_pca_batch_after.png": "Figure 15. PCA visualization of batch effects after normalization.",
    "figure16_feature_ablation.png": "Figure 16. Feature-ablation impact on Random Forest ROC AUC.",
}

TABLE_SPECS = [
    ("results_summary/advanced_model_metrics.csv", "Table 1. Model performance metrics with 95% bootstrap confidence intervals."),
    ("results_summary/model_comparison_transfer_learning.csv", "Table 2. Transfer learning comparison: CNN-scratch vs ResNet-scratch vs ResNet-pretrained."),
    ("results_summary/advanced_calibration_ece.csv", "Table 2. Calibration error summary (ECE)."),
    ("results_summary/advanced_feature_ablation.csv", "Table 3. Feature-ablation outcomes."),
    ("results_summary/advanced_biological_validation.csv", "Table 4. Biological validation statistics for top features."),
    ("results_summary/advanced_computational_cost.csv", "Table 5. Computational cost comparison (train time, peak RAM, pretrained weights)."),
    ("results_summary/advanced_nested_cv.csv", "Table 6. Nested cross-validation stability summary."),
    ("results_summary/robustness_classification.csv", "Table 7. Robustness classification across threshold settings."),
]


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_normal_font(doc: Document) -> None:
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(12)


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_dataframe_table(doc: Document, df: pd.DataFrame, caption: str) -> None:
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_json_table(doc: Document, data: dict, caption: str) -> None:
    rows = []
    for cmp_name, values in data.items():
        if isinstance(values, dict):
            row = {"comparison": cmp_name}
            row.update(values)
            rows.append(row)
    if rows:
        add_dataframe_table(doc, pd.DataFrame(rows), caption)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.2))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_markdown_content(doc: Document, md_text: str) -> None:
    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(16)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. ") or line.startswith("7. ") or line.startswith("8. ") or line.startswith("9. "):
            doc.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def main() -> None:
    if TEMPLATE_PATH.exists():
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    clear_document(doc)
    set_normal_font(doc)

    md = MARKDOWN_PATH.read_text(encoding="utf-8")
    add_markdown_content(doc, md)

    doc.add_heading("Figures", level=1)
    fig_dir = ROOT / "final_figures"
    for fig_name in sorted(FIGURE_CAPTIONS.keys(), key=lambda x: [int(s) if s.isdigit() else s for s in x.replace('.png','').replace('figure','').replace('_',' ').split()]):
        fig_path = fig_dir / fig_name
        if fig_path.exists():
            add_figure(doc, fig_path, FIGURE_CAPTIONS[fig_name])

    doc.add_heading("Tables", level=1)
    for rel, caption in TABLE_SPECS:
        p = ROOT / rel
        if p.exists():
            add_dataframe_table(doc, pd.read_csv(p), caption)

    delong = ROOT / "results_summary" / "advanced_delong_tests.json"
    if delong.exists():
        add_json_table(doc, json.loads(delong.read_text(encoding="utf-8")), "Table 8. DeLong test results comparing deep-learning AUC against Random Forest.")

    doc.add_heading("Figure Legends", level=1)
    for k in sorted(FIGURE_CAPTIONS.keys()):
        doc.add_paragraph(FIGURE_CAPTIONS[k])

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
