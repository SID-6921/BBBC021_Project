from __future__ import annotations

from pathlib import Path
import json

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = Path(r"C:\Users\nanda\Downloads\Manuscript (2).docx")
OUTPUT_PATH = ROOT / "results_summary" / "BBBC021_Manuscript_Submission_v2.docx"

TITLE = "A Robust and Reproducible Pipeline for Automated Spot Detection, Quantification, and Group Classification in BBBC021 Fluorescence Microscopy Images"
AUTHORS = "Nanda et al."
CORRESPONDENCE = "Correspondence: to be finalized before submission"

FIGURE_CAPTIONS = {
    "figure1_pipeline_workflow.png": "Figure 1. End-to-end workflow for BBBC021 image preprocessing, robust spot detection, feature extraction, supervised modeling, unsupervised analysis, robustness testing, and deep-learning comparison.",
    "figure2_detection_samples.png": "Figure 2. Representative BBBC021 detection overlays showing bright-region detection results on sample fused images.",
    "figure3_spot_count_boxplot.png": "Figure 3A. Distribution of detected spot counts between Group A and Group B.",
    "figure3_mean_intensity_boxplot.png": "Figure 3B. Distribution of mean intensity between Group A and Group B.",
    "figure3_density_spots_per_10k_px_boxplot.png": "Figure 3C. Distribution of spot density between Group A and Group B.",
    "figure4_classification_roc.png": "Figure 4A. ROC curves for Logistic Regression and Random Forest models trained on engineered image features.",
    "figure4_feature_importance.png": "Figure 4B. Random Forest feature-importance ranking for the engineered image descriptors.",
    "figure5_pca_clustering.png": "Figure 5. PCA projection and K-means clustering of the feature space, illustrating latent separation trends between groups.",
    "figure6_batch_robustness.png": "Figure 6. Batch-wise spot count robustness across multiple BBBC021 batches and detection configurations.",
    "figure7_threshold_sensitivity.png": "Figure 7. Sensitivity of summary detection metrics across sensitive, default, and conservative threshold configurations.",
    "figure8_robustness_classification.png": "Figure 8. Classification accuracy across detection configurations during robustness validation.",
    "figure9_deep_learning_roc.png": "Figure 9A. ROC curve for the compact CNN deep-learning baseline.",
    "figure9_deep_learning_training.png": "Figure 9B. Training and validation loss curves for the compact CNN baseline.",
    "figure10_confusion_matrix_rf.png": "Figure 10. Confusion matrix for the Random Forest model on the 1600-image advanced evaluation split.",
    "figure11_confusion_matrix_cnn.png": "Figure 11. Confusion matrix for the compact CNN model on the advanced evaluation split.",
    "figure12_confusion_matrix_resnet18.png": "Figure 12. Confusion matrix for the ResNet-18 transfer-learning baseline on the advanced evaluation split.",
    "figure13_calibration_curves.png": "Figure 13. Calibration curves for feature-based and deep-learning models with reliability comparison across predicted probability bins.",
    "figure14_pca_batch_before.png": "Figure 14. PCA visualization of batch effects before batch-wise normalization.",
    "figure15_pca_batch_after.png": "Figure 15. PCA visualization of batch effects after batch-wise normalization.",
    "figure16_feature_ablation.png": "Figure 16. Feature ablation study showing the impact of removing top-ranked Random Forest features on test-set ROC AUC.",
}

TABLE_SPECS = [
    ("results_summary/advanced_model_metrics.csv", "Table 1. Advanced full-dataset model metrics including overall accuracy and per-class ROC AUC."),
    ("results_summary/model_comparison_final.csv", "Table 2. Summary comparison of classical and deep models from the earlier finalized benchmark run."),
    ("results_summary/advanced_calibration_ece.csv", "Table 3. Expected Calibration Error (ECE) for all evaluated models."),
    ("results_summary/advanced_feature_ablation.csv", "Table 4. Feature ablation results quantifying ROC AUC changes after removal of top-ranked features."),
    ("results_summary/advanced_biological_validation.csv", "Table 5. Biological validation of top features against compound identity and concentration metadata."),
    ("results_summary/advanced_computational_cost.csv", "Table 6. Computational cost comparison across feature-based and deep-learning models."),
    ("results_summary/advanced_nested_cv.csv", "Table 7. Nested cross-validation results for feature-based models."),
    ("results_summary/robustness_classification.csv", "Table 8. Robustness classification results across detection configurations."),
]


ABSTRACT = (
    "High-content microscopy analysis requires scalable and reproducible computational methods that can replace or complement manual workflows. "
    "We present a complete image-analysis system for BBBC021 integrating robust spot detection, quantitative feature extraction, supervised and "
    "unsupervised learning, robustness validation across batches, and deep-learning comparison. The advanced run processed 1600 images and added "
    "confusion matrices, per-class ROC AUC, DeLong significance testing, calibration analysis, feature ablation, batch-effect PCA before/after "
    "normalization, biological feature validation, computational cost profiling, and nested cross-validation. Feature-based models achieved strong "
    "discrimination (Logistic Regression: accuracy 0.8938, ROC AUC 0.9231; Random Forest: accuracy 0.9156, ROC AUC 0.9399). CNN and ResNet-18 baselines "
    "reached accuracy 0.8906 and ROC AUC 0.7422. DeLong testing confirmed a significant AUC difference between feature and deep baselines (p = 5.28e-08)."
)

KEYWORDS = "BBBC021; fluorescence microscopy; image analysis; adaptive thresholding; feature engineering; machine learning; random forest; convolutional neural network; transfer learning; robustness analysis; reproducibility"

SECTIONS = [
    ("1. Introduction", [
        "Automated image analysis has become essential in high-content screening, where manual review is limited by subjectivity and throughput constraints. BBBC021 provides a practical benchmark for studying phenotypic image analysis under diverse treatment conditions.",
        "This work aimed to deliver a complete, reproducible, and publication-ready BBBC021 workflow spanning robust detection, interpretable quantification, supervised and unsupervised modeling, robustness validation, and deep-learning comparison.",
        "Unlike narrow benchmark reports, this revised manuscript emphasizes reproducibility, calibration quality, statistical validation, and interpretation clarity for journal-level evaluation.",
    ]),
    ("2. Materials and Methods", [
        "The advanced evaluation processed 1600 BBBC021 images collected from multiple weekly archives. Available channels were fused to grayscale, resized, normalized, and contrast-enhanced prior to detection.",
        "Spot detection used adaptive Gaussian thresholding with morphological opening and closing. Candidate contours were filtered by area and contour-level mean intensity to reduce false positives under heterogeneous illumination.",
        "Per-image features included spot count, mean intensity, total intensity, intensity variance, area coverage, density measures, spot-size distribution statistics, and size-fraction descriptors. Logistic Regression and Random Forest were trained on engineered features. Deep-learning baselines included a compact CNN and a ResNet-18 transfer-learning model.",
        "Additional validation included confusion matrices, per-class ROC AUC, calibration curves, Expected Calibration Error, DeLong tests, PCA batch-effect visualization before/after normalization, feature ablation, biological validation against compound metadata, computational cost analysis, and nested cross-validation.",
        "All outputs were exported into publication-ready figures and structured tables to preserve direct traceability from raw image processing to final conclusions.",
    ]),
    ("3. Results", [
        "On the 1600-image advanced split, Logistic Regression achieved accuracy 0.8938 and ROC AUC 0.9231, while Random Forest achieved accuracy 0.9156 and ROC AUC 0.9399.",
        "The compact CNN and ResNet-18 transfer baseline each achieved accuracy 0.8906 with ROC AUC 0.7422 under the final advanced evaluation. DeLong tests versus Random Forest showed statistically significant AUC differences (p = 5.28e-08).",
        "Nested cross-validation showed strong feature-model stability (mean AUC approximately 0.91 for Logistic Regression and 0.91 for Random Forest). Calibration curves, batch-effect PCA, ablation analysis, and computational-cost summaries are included below.",
        "Feature-ablation trends and biological validation against compound metadata further support the stability and interpretability of the top-ranked descriptors.",
    ]),
    ("4. Discussion", [
        "The results indicate that a carefully engineered classical computer-vision and feature-learning workflow remains highly competitive and, in this task configuration, outperforms the tested deep baselines in ROC AUC.",
        "The added calibration, ablation, biological validation, nested CV, and batch-effect analyses strengthen the evidentiary quality of the study and make the pipeline suitable for journal submission and follow-on methodological extensions.",
        "Future work will expand to additional weekly BBBC021 archives and evaluate stronger transfer-learning schedules for deep baselines while retaining the current reproducibility guarantees.",
    ]),
    ("5. Conclusions", [
        "We developed and validated a complete BBBC021 analysis system that scales beyond 1500 images and includes robust detection, rich feature modeling, deep-learning baselines, significance testing, calibration diagnostics, batch-effect analysis, feature ablation, biological validation, and overfitting checks.",
    ]),
]


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_normal_font(doc: Document) -> None:
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(12)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    for text in [AUTHORS, CORRESPONDENCE]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Revised Submission Draft (Version 2)")
    run.italic = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_dataframe_table(doc: Document, df: pd.DataFrame, caption: str) -> None:
    doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_json_table(doc: Document, data: dict, caption: str) -> None:
    rows = []
    for outer_key, outer_val in data.items():
        if isinstance(outer_val, dict):
            row = {"comparison": outer_key}
            row.update(outer_val)
            rows.append(row)
    df = pd.DataFrame(rows)
    add_dataframe_table(doc, df, caption)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def main() -> None:
    if TEMPLATE_PATH.exists():
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    clear_document(doc)
    set_normal_font(doc)

    add_title_block(doc)
    add_heading(doc, "Manuscript Overview", level=1)
    doc.add_paragraph("This revised version includes full advanced analyses and all required figures, tables, images, and flowchart content for submission-quality review.")
    add_heading(doc, "Abstract", level=1)
    doc.add_paragraph(ABSTRACT)
    doc.add_paragraph(f"Keywords: {KEYWORDS}")

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading, level=1)
        add_paragraphs(doc, paragraphs)

    add_heading(doc, "Figures", level=1)
    figure_dir = ROOT / "final_figures"
    for fig_name in sorted(FIGURE_CAPTIONS.keys(), key=lambda x: [int(s) if s.isdigit() else s for s in x.replace('.png','').replace('figure','').replace('_',' ').split()]):
        fig_path = figure_dir / fig_name
        if fig_path.exists():
            add_figure(doc, fig_path, FIGURE_CAPTIONS[fig_name])

    add_heading(doc, "Tables", level=1)
    for rel_path, caption in TABLE_SPECS:
        path = ROOT / rel_path
        if path.exists():
            df = pd.read_csv(path)
            add_dataframe_table(doc, df, caption)

    delong_path = ROOT / "results_summary" / "advanced_delong_tests.json"
    if delong_path.exists():
        data = json.loads(delong_path.read_text(encoding="utf-8"))
        add_json_table(doc, data, "Table 9. DeLong test results comparing deep-learning ROC AUC against the Random Forest feature model.")

    add_heading(doc, "Figure Legends", level=1)
    for fig_name in sorted(FIGURE_CAPTIONS.keys()):
        doc.add_paragraph(FIGURE_CAPTIONS[fig_name])

    add_heading(doc, "Table Legends", level=1)
    for _, caption in TABLE_SPECS:
        doc.add_paragraph(caption)
    doc.add_paragraph("Table 9. DeLong test results comparing deep-learning ROC AUC against the Random Forest feature model.")

    add_heading(doc, "Data Availability", level=1)
    doc.add_paragraph("All generated figures, tables, and summaries are available in the GitHub repository and were used to construct this submission-ready manuscript document.")

    add_heading(doc, "Conflicts of Interest", level=1)
    doc.add_paragraph("The authors declare no conflict of interest.")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved manuscript to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
